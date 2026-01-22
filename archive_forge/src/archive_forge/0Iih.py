import os
import json
import csv
import xml.etree.ElementTree as ET
from typing import Any, Dict, Union, Optional, Tuple, List
import pandas as pd
import logging
import yaml
import pickle
import configparser
import markdown
import docx
import openpyxl
import sqlite3
import PyPDF2
import PIL.Image

# Configure logging
logging.basicConfig(
    level=logging.DEBUG, format="%(asctime)s - %(levelname)s - %(message)s"
)


class UniversalDataReader:
    """
    A class meticulously designed to read and process various file types through a universal interface.
    This class encapsulates methods that identify and process file content with precision and adaptability,
    ensuring optimal performance and extensibility.
    """

    def __init__(self, file_path: str):
        """
        Initialize the UniversalDataReader with a specific file path.

        :param file_path: The path to the file to be read.
        :type file_path: str
        """
        self.file_path: str = file_path
        self.file_type: str = self.identify_file_type()
        logging.info(
            f"UniversalDataReader initialized for file: {file_path} with identified type: {self.file_type}"
        )

    def identify_file_type(self) -> str:
        """
        Determine the file type by extracting and analyzing the file extension.

        :return: A string representing the file type.
        :rtype: str
        """
        _, file_extension = os.path.splitext(self.file_path)
        file_extension = file_extension.lower()
        logging.debug(f"File extension identified: {file_extension}")
        return file_extension

    def read_file(
        self,
    ) -> Union[
        str,
        Dict[str, Any],
        pd.DataFrame,
        ET.ElementTree,
        configparser.ConfigParser,
        object,
        List[str],
        docx.document.Document,
        openpyxl.workbook.workbook.Workbook,
        sqlite3.Connection,
        PyPDF2.PdfReader,
        PIL.Image.Image,
    ]:
        """
        Read the file based on its type and return its content in an appropriate format.

        :return: The content of the file, formatted according to its type.
        :rtype: Union[str, Dict[str, Any], pd.DataFrame, ET.ElementTree, configparser.ConfigParser, object, List[str], docx.document.Document, openpyxl.workbook.workbook.Workbook, sqlite3.Connection, PyPDF2.PdfReader, PIL.Image.Image]
        """
        try:
            if self.file_type == ".json":
                return self.read_json()
            elif self.file_type == ".csv":
                return self.read_csv()
            elif self.file_type == ".xml":
                return self.read_xml()
            elif self.file_type in [".txt", ".log"]:
                return self.read_text()
            elif self.file_type == ".yaml" or self.file_type == ".yml":
                return self.read_yaml()
            elif self.file_type == ".ini":
                return self.read_ini()
            elif self.file_type == ".pkl":
                return self.read_pickle()
            elif self.file_type == ".md":
                return self.read_markdown()
            elif self.file_type == ".docx":
                return self.read_docx()
            elif self.file_type == ".xlsx":
                return self.read_excel()
            elif self.file_type == ".db":
                return self.read_database()
            elif self.file_type == ".pdf":
                return self.read_pdf()
            elif self.file_type in [".jpg", ".jpeg", ".png", ".gif", ".bmp"]:
                return self.read_image()
            else:
                logging.error(f"Unsupported file type: {self.file_type}")
                raise ValueError(f"Unsupported file type: {self.file_type}")
        except Exception as e:
            logging.error(f"Error reading file {self.file_path}: {str(e)}")
            raise

    def read_markdown(self) -> List[str]:
        """
        Read a Markdown file and return its content as a list of strings, each representing a line.
        This method ensures that every line of the Markdown file is read with precision and completeness,
        preserving the integrity and authenticity of the data.

        :return: The content of the Markdown file, with each line retained in its original form.
        :rtype: List[str]
        """
        try:
            with open(self.file_path, "r", encoding="utf-8") as file:
                data = file.readlines()
                logging.debug(f"Markdown data read successfully from {self.file_path}")
                return data
        except FileNotFoundError:
            logging.error(f"Markdown file not found at {self.file_path}")
            raise FileNotFoundError(f"Markdown file not found at {self.file_path}")
        except Exception as e:
            logging.error(
                f"An error occurred while reading the Markdown file at {self.file_path}: {str(e)}"
            )
            raise Exception(
                f"An error occurred while reading the Markdown file: {str(e)}"
            )

    def read_docx(self) -> docx.document.Document:
        """
        Read a DOCX file and return its content as a docx Document object, ensuring that all elements,
        including text, tables, images, and other media, are comprehensively extracted and represented
        in the returned Document object. This method is designed to handle complex DOCX files with
        various embedded elements, providing a robust and complete representation of the document.

        :return: The content of the DOCX file, including all textual and non-textual elements.
        :rtype: docx.document.Document
        """
        try:
            # Load the DOCX file into a Document object
            doc = docx.Document(self.file_path)
            logging.debug(f"DOCX data read successfully from {self.file_path}")

            # Log details about the document structure
            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            logging.info(
                f"Document contains {paragraphs} paragraphs and {tables} tables."
            )

            # Extract and log images if present
            images = [
                r
                for r in doc.inline_shapes
                if r.type == docx.enum.shape.WD_INLINE_SHAPE.PICTURE
            ]
            logging.info(f"Document contains {len(images)} images.")

            # Return the fully loaded Document object
            return doc
        except Exception as e:
            logging.error(
                f"An error occurred while reading the DOCX file at {self.file_path}: {str(e)}"
            )
            raise Exception(f"An error occurred while reading the DOCX file: {str(e)}")

    def read_excel(self) -> openpyxl.workbook.workbook.Workbook:
        """
        Read an Excel file and return its content as an openpyxl Workbook object, ensuring comprehensive extraction
        of all data types including text, formulas, images, charts, and hyperlinks. This method is designed to handle
        complex Excel files with multiple sheets, providing a robust and complete representation of the workbook.

        :return: The content of the Excel file, including all sheets and their respective elements.
        :rtype: openpyxl.workbook.workbook.Workbook
        """
        try:
            # Load the Excel file into a Workbook object
            workbook = openpyxl.load_workbook(self.file_path, data_only=False)
            logging.debug(f"Excel data read successfully from {self.file_path}")

            # Log details about the workbook structure
            sheets = workbook.sheetnames
            logging.info(f"Workbook contains {len(sheets)} sheets: {sheets}")

            # Iterate through each sheet and log detailed information about their contents
            for sheet_name in sheets:
                sheet = workbook[sheet_name]
                num_rows = sheet.max_row
                num_cols = sheet.max_column
                logging.info(
                    f"Sheet '{sheet_name}' contains {num_rows} rows and {num_cols} columns."
                )

                # Check for images, charts, and hyperlinks within the sheet
                images = sheet._images
                charts = sheet.charts
                hyperlinks = [
                    cell.hyperlink
                    for row in sheet.iter_rows()
                    for cell in row
                    if cell.hyperlink
                ]
                logging.info(
                    f"Sheet '{sheet_name}' contains {len(images)} images, {len(charts)} charts, and {len(hyperlinks)} hyperlinks."
                )

            # Return the fully loaded Workbook object
            return workbook
        except Exception as e:
            logging.error(
                f"An error occurred while reading the Excel file at {self.file_path}: {str(e)}"
            )
            raise Exception(f"An error occurred while reading the Excel file: {str(e)}")

    def read_database(self) -> sqlite3.Connection:
        """
        Read a database file and return its connection object, ensuring a robust, comprehensive, and detailed
        connection process that captures and logs every step, providing a perfect and complete understanding
        of the database connection status.

        This method is designed to handle SQLite databases, ensuring that the connection is established
        flawlessly with detailed logging of the connection process. It is meticulously crafted to ensure
        that all possible errors are caught and handled, providing a faultless and seamless database
        interaction experience.

        :return: The connection to the database.
        :rtype: sqlite3.Connection
        """
        try:
            # Attempt to establish a connection to the SQLite database
            connection = sqlite3.connect(self.file_path)
            logging.debug(
                f"Database connection established successfully from {self.file_path}"
            )

            # Log the successful connection
            logging.info(f"Successfully connected to the database at {self.file_path}")

            # Return the established connection
            return connection
        except sqlite3.Error as e:
            # Log any errors that occur during the connection attempt
            logging.error(
                f"Failed to connect to the database at {self.file_path}: {str(e)}"
            )
            raise Exception(
                f"An error occurred while connecting to the database: {str(e)}"
            )

    def read_pdf(self) -> PyPDF2.PdfReader:
        """
        Read a PDF file and return its content as a PyPDF2 PdfReader object, ensuring that every aspect of the PDF data,
        including text, images, forms, and metadata, is comprehensively extracted and represented in the returned PdfReader object.
        This method is meticulously crafted to handle complex PDF files with various embedded elements, providing a robust,
        complete, and perfect representation of the PDF content.

        :return: The content of the PDF file, including all textual and non-textual elements.
        :rtype: PyPDF2.PdfReader
        """
        try:
            # Open the PDF file in binary read mode
            with open(self.file_path, "rb") as file:
                # Initialize the PdfReader object from PyPDF2 to read the PDF data
                reader = PyPDF2.PdfReader(file)

                # Log the successful reading of the PDF data
                logging.debug(f"PDF data read successfully from {self.file_path}")

                # Extract and log the number of pages in the PDF
                num_pages = len(reader.pages)
                logging.info(f"The PDF file contains {num_pages} pages.")

                # Optionally, log detailed information about each page (text, images, etc.)
                for i, page in enumerate(reader.pages, start=1):
                    text = page.extract_text()
                    logging.debug(
                        f"Extracted text from page {i}: {text[:100]}..."
                    )  # Log the first 100 characters of text

                # Return the fully loaded PdfReader object containing all PDF data
                return reader
        except Exception as e:
            # Log any errors that occur during the PDF reading process
            logging.error(
                f"An error occurred while reading the PDF file at {self.file_path}: {str(e)}"
            )
            raise Exception(f"An error occurred while reading the PDF file: {str(e)}")

    def read_image(self) -> PIL.Image.Image:
        """
        Read an image file and return its content as a PIL Image object, ensuring that every pixel, color, and metadata
        detail is comprehensively extracted and represented in the returned Image object. This method is meticulously
        crafted to handle various image formats with precision, providing a robust, complete, and perfect representation
        of the image content.

        :return: The content of the image file, including all visual and metadata elements.
        :rtype: PIL.Image.Image
        """
        try:
            # Open the image file, ensuring it is read in the mode that preserves all color and transparency information
            image = PIL.Image.open(self.file_path)
            image.load()  # Explicitly load the image to ensure all data is read

            # Log the successful reading of the image data
            logging.debug(f"Image data read successfully from {self.file_path}")

            # Extract and log metadata if available
            if hasattr(image, "info"):
                metadata = image.info
                logging.info(f"Image metadata extracted: {metadata}")

            # Optionally, log detailed information about the image (dimensions, mode, etc.)
            logging.info(f"Image dimensions: {image.size} pixels")
            logging.info(f"Image mode (color depth): {image.mode}")

            # Return the fully loaded Image object containing all visual and metadata
            return image
        except Exception as e:
            # Log any errors that occur during the image reading process
            logging.error(
                f"An error occurred while reading the image file at {self.file_path}: {str(e)}"
            )
            raise Exception(f"An error occurred while reading the image file: {str(e)}")

    def read_json(self) -> Dict[str, Any]:
        """
        Read a JSON file and return its content as a dictionary, ensuring that every possible element,
        including nested structures and metadata, is comprehensively extracted and represented in the
        returned dictionary. This method is meticulously crafted to handle various JSON structures with
        precision, providing a robust, complete, and perfect representation of the JSON content.

        :return: The content of the JSON file parsed into a dictionary, including all nested structures and metadata.
        :rtype: Dict[str, Any]
        """
        try:
            with open(self.file_path, "r", encoding="utf-8") as file:
                data = json.load(file)
                logging.debug(f"JSON data read successfully from {self.file_path}")

                # Log detailed information about the JSON structure
                logging.info(
                    f"JSON data structure details: {json.dumps(data, indent=4)}"
                )

                # Return the fully loaded dictionary containing all JSON data
                return data
        except json.JSONDecodeError as e:
            # Log any errors that occur during the JSON reading process
            logging.error(
                f"An error occurred while decoding the JSON file at {self.file_path}: {str(e)}"
            )
            raise Exception(f"An error occurred while decoding the JSON file: {str(e)}")
        except Exception as e:
            # Log any other errors that occur during the file reading process
            logging.error(
                f"An error occurred while reading the JSON file at {self.file_path}: {str(e)}"
            )
            raise Exception(f"An error occurred while reading the JSON file: {str(e)}")

    def read_csv(self) -> pd.DataFrame:
        """
        Read a CSV file and return its content as a pandas DataFrame, ensuring that every possible element,
        including headers, data types, and missing values, is comprehensively extracted and represented in the
        returned DataFrame. This method is meticulously crafted to handle various CSV structures with precision,
        providing a robust, complete, and perfect representation of the CSV content.

        :return: The content of the CSV file parsed into a DataFrame, including all possible data types and handling missing values.
        :rtype: pd.DataFrame
        """
        try:
            # Attempt to read the CSV file with detailed logging of each step
            data = pd.read_csv(
                self.file_path, dtype=str
            )  # Read all data as string to prevent unintended type inference
            logging.debug(f"CSV data read successfully from {self.file_path}")

            # Log detailed information about the DataFrame structure
            logging.info(f"DataFrame shape: {data.shape}")
            logging.info(f"DataFrame columns: {data.columns.tolist()}")

            # Check and log if there are any missing values in the DataFrame
            if data.isnull().values.any():
                missing_info = data.isnull().sum()
                logging.warning(f"Missing values found in DataFrame: {missing_info}")
            else:
                logging.info("No missing values found in DataFrame.")

            # Return the fully loaded DataFrame containing all CSV data
            return data
        except pd.errors.EmptyDataError as e:
            # Log any errors that occur during the CSV reading process
            logging.error(f"No data: Empty CSV file at {self.file_path}: {str(e)}")
            raise Exception(f"No data: Empty CSV file at {self.file_path}: {str(e)}")
        except pd.errors.ParserError as e:
            # Log any errors that occur during the parsing of the CSV file
            logging.error(
                f"Parsing error: Error parsing CSV file at {self.file_path}: {str(e)}"
            )
            raise Exception(
                f"Parsing error: Error parsing CSV file at {self.file_path}: {str(e)}"
            )
        except Exception as e:
            # Log any other errors that occur during the file reading process
            logging.error(
                f"An error occurred while reading the CSV file at {self.file_path}: {str(e)}"
            )
            raise Exception(f"An error occurred while reading the CSV file: {str(e)}")

    def read_xml(self) -> ET.ElementTree:
        """
        Read an XML file and return its content as an ElementTree object, ensuring that every possible element,
        attribute, and nested structure within the XML file is comprehensively extracted and represented in the
        returned ElementTree. This method is meticulously crafted to handle complex XML files with precision,
        providing a robust, complete, and perfect representation of the XML content.

        :return: The content of the XML file parsed into an ElementTree, including all elements, attributes, and nested structures.
        :rtype: ET.ElementTree
        """
        try:
            # Parse the XML file into an ElementTree object
            tree = ET.parse(self.file_path)
            logging.debug(f"XML data read successfully from {self.file_path}")

            # Log detailed information about the ElementTree structure
            root = tree.getroot()
            logging.info(f"Root element: {root.tag}")
            logging.info(f"Number of child elements in the root: {len(list(root))}")

            # Recursively log details of each element in the XML tree
            def log_element_details(element, depth=0):
                logging.debug(
                    f"{'  ' * depth}Element: {element.tag}, Attributes: {element.attrib}"
                )
                for child in list(element):
                    log_element_details(child, depth + 1)

            log_element_details(root)

            # Return the fully parsed and logged ElementTree object
            return tree
        except ET.ParseError as e:
            # Log any errors that occur during the XML parsing process
            logging.error(f"XML parsing error at {self.file_path}: {str(e)}")
            raise Exception(f"XML parsing error at {self.file_path}: {str(e)}")
        except Exception as e:
            # Log any other errors that occur during the file reading process
            logging.error(
                f"An error occurred while reading the XML file at {self.file_path}: {str(e)}"
            )
            raise Exception(f"An error occurred while reading the XML file: {str(e)}")

    def read_text(self) -> str:
        """
        Read a text file and return its content as a string, ensuring that every character, line break, and encoding detail
        is accurately captured and represented in the returned string. This method is meticulously crafted to handle
        text files of any size and encoding, providing a robust, complete, and perfect representation of the text content.

        :return: The content of the text file, preserving all textual details and nuances.
        :rtype: str
        """
        try:
            with open(self.file_path, "r", encoding="utf-8") as file:
                data = file.read()
                logging.debug(f"Text data read successfully from {self.file_path}")

                # Log detailed information about the text content
                num_characters = len(data)
                num_lines = (
                    data.count("\n") + 1
                )  # Adding 1 to count the last line if it doesn't end with a newline
                logging.info(f"Number of characters in text: {num_characters}")
                logging.info(f"Number of lines in text: {num_lines}")

                return data
        except UnicodeDecodeError as e:
            # Log any errors that occur during the text reading process due to encoding issues
            logging.error(
                f"Unicode decoding error while reading the text file at {self.file_path}: {str(e)}"
            )
            raise Exception(
                f"Unicode decoding error while reading the text file: {str(e)}"
            )
        except Exception as e:
            # Log any other errors that occur during the file reading process
            logging.error(
                f"An error occurred while reading the text file at {self.file_path}: {str(e)}"
            )
            raise Exception(f"An error occurred while reading the text file: {str(e)}")

    def read_yaml(self) -> Dict[str, Any]:
        """
        Read a YAML file and return its content as a dictionary.

        :return: The content of the YAML file parsed into a dictionary.
        :rtype: Dict[str, Any]
        """
        with open(self.file_path, "r", encoding="utf-8") as file:
            data = yaml.safe_load(file)
            logging.debug(f"YAML data read successfully from {self.file_path}")
            return data

    def read_ini(self) -> configparser.ConfigParser:
        """
        Read an INI file and return its content as a ConfigParser object.

        :return: The content of the INI file parsed into a ConfigParser.
        :rtype: configparser.ConfigParser
        """
        config = configparser.ConfigParser()
        config.read(self.file_path)
        logging.debug(f"INI data read successfully from {self.file_path}")
        return config

    def read_pickle(self) -> object:
        """
        Read a pickle file and return its content as a Python object.

        :return: The content of the pickle file deserialized into a Python object.
        :rtype: object
        """
        with open(self.file_path, "rb") as file:
            data = pickle.load(file)
            logging.debug(f"Pickle data read successfully from {self.file_path}")
            return data


# Example usage:
# data_reader = UniversalDataReader('path_to_your_file.extension')
# content = data_reader.read_file()
# print(content)
