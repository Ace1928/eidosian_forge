import asyncio  # For asynchronous programming, allowing non-blocking I/O operations.
import logging  # For logging events and debugging information.
import nest_asyncio  # For applying monkey-patching to allow nested asyncio loops.

nest_asyncio.apply()  # Apply the patch to enable nested asyncio loops.

from typing import Any, List, Callable, Optional, Union, Dict, Tuple, Sequence # For type hinting and annotations.
from IPython.display import Markdown, display # For displaying markdown and other rich content in IPython environments.

from llama_index.core.async_utils import run_jobs # For running asynchronous jobs in LlamaIndex.
from llama_index.core.indices.property_graph.utils import ( # For utility functions related to property graphs in LlamaIndex.
    default_parse_triplets_fn,
)
from llama_index.core.graph_stores.types import ( # For type definitions related to graph stores in LlamaIndex.
    EntityNode,
    KG_NODES_KEY,
    KG_RELATIONS_KEY,
    Relation,
)
from llama_index.core.llms.llm import LLM # For the base class of Large Language Models in LlamaIndex.
from llama_index.core.prompts import PromptTemplate # For creating and managing prompts in LlamaIndex.
from llama_index.core.prompts.default_prompts import ( # For default prompt templates in LlamaIndex.
    DEFAULT_KG_TRIPLET_EXTRACT_PROMPT,
)
from llama_index.core.schema import TransformComponent, BaseNode # For schema definitions in LlamaIndex.
from llama_index.core.bridge.pydantic import BaseModel, Field # For data validation and model creation using Pydantic in LlamaIndex.

import os # For interacting with the operating system, such as file paths.
import re # For regular expression operations.
import zipfile # For working with zip archives.
import pickle # For serializing and deserializing Python objects.
import hashlib # For creating hash values.
import traceback # For printing or retrieving stack traces.
import multiprocessing # For creating and managing processes.
from concurrent.futures import ThreadPoolExecutor, as_completed # For managing concurrent tasks using threads.
from io import BytesIO # For working with in-memory binary streams.

import pandas as pd # For data manipulation and analysis using DataFrames.
import torch # For tensor operations and deep learning.
from tqdm.auto import tqdm # For displaying progress bars.
from PyPDF2 import PdfReader # For reading PDF files.
import docx # For reading docx files.
import networkx as nx # For creating and manipulating graphs.
from graspologic.partition import hierarchical_leiden # For graph partitioning algorithms.

from accelerate import disk_offload # For offloading model weights to disk.
from transformers import AutoModelForCausalLM, AutoTokenizer # For loading and using pre-trained transformer models.

from llama_index.core import Document, Settings, PropertyGraphIndex # For core LlamaIndex classes.
from llama_index.core.chat_engine.types import ( # For type definitions related to chat engines in LlamaIndex.
    ChatResponseMode,
    ChatMode,
    AgentChatResponse,
)
from llama_index.core.graph_stores.types import ( # For type definitions related to graph stores in LlamaIndex.
    PropertyGraphStore,
)
from llama_index.core.graph_stores import SimplePropertyGraphStore # For a simple implementation of a property graph store in LlamaIndex.
from llama_index.core.llms import ChatMessage, MessageRole # For chat message types in LlamaIndex.
from llama_index.core.node_parser import SentenceSplitter # For splitting text into sentences in LlamaIndex.
from llama_index.core.query_engine import CustomQueryEngine # For creating custom query engines in LlamaIndex.
from llama_index.core.schema import MetadataMode # For metadata handling in LlamaIndex.
from llama_index.embeddings.huggingface import HuggingFaceEmbedding # For using Hugging Face models for embeddings in LlamaIndex.
from llama_index.llms.huggingface import HuggingFaceLLM # For using Hugging Face models as LLMs in LlamaIndex.
from networkx import graph, Graph # For graph data structures.
from huggingface_hub import hf_hub_download # For downloading files from the Hugging Face Hub.
from eidos_config import DEFAULT_HF_TOKEN # For default Hugging Face token.


# Configure logging to display messages in real-time for debugging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
# Configure logging
logger = logging.getLogger(__name__)


# Define the default model name for Qwen
MODEL_NAME = "Qwen/Qwen2.5-0.5B-Instruct"
# Define the local directory to save the model
SAVE_DIR = "./saved_models"
# Construct the full path where the model will be saved
MODEL_SAVE_PATH = os.path.join(SAVE_DIR, MODEL_NAME)

# Ensure the save directory exists; create it if it doesn't
os.makedirs(SAVE_DIR, exist_ok=True)

# Check if the model already exists at the specified path
if not os.path.exists(MODEL_SAVE_PATH):
    # Log the initiation of the model download process
    logging.info(f"Downloading {MODEL_NAME} from Hugging Face Hub...")
    try:
        # Attempt to download the model from Hugging Face Hub
        hf_hub_download(
            repo_id=str(MODEL_NAME),  # The model's repository ID on Hugging Face Hub
            local_dir=SAVE_DIR,  # The local directory to save the model
            filename=MODEL_NAME,  # The specific model file to download
            cache_dir="qwen_offload_cache",  # Cache directory for downloaded files
            token=DEFAULT_HF_TOKEN,
        )
        # Log a success message upon successful download
        logging.info(f"Successfully downloaded {MODEL_NAME} to {MODEL_SAVE_PATH}")
    except Exception as e:
        # Log an error message if the download fails, including the exception details
        logging.error(f"Failed to download {MODEL_NAME}: {e}")
else:
    # Log a message indicating that the model already exists locally
    logging.info(
        f"{MODEL_NAME} already exists at {MODEL_SAVE_PATH}. Skipping download."
    )


def build_qwen_llm(
    model_name: str = MODEL_NAME,
    max_new_tokens: int = 128,
    device_map: str = "auto",
    trust_remote_code: bool = True,
    dtype="float32",
) -> HuggingFaceLLM:
    """
    Build a HuggingFaceLLM that loads Qwen (or any other Hugging Face model)
    and ensures correct parameters for generation.

    This function initializes a HuggingFaceLLM instance with specific configurations
    tailored for the Qwen model, including setting the maximum number of new tokens,
    device mapping, and trust for remote code.

    Args:
        model_name (str): The Hugging Face model hub path for the Qwen model.
            Defaults to the global `MODEL_NAME`.
        max_new_tokens (int): The maximum number of tokens to generate.
            Defaults to 4096.
        device_map (str): The device map for loading the model ('auto', 'cpu', 'cuda', etc.).
            Defaults to 'auto'.
        trust_remote_code (bool): Whether to trust remote code from the Hugging Face model hub.
            Defaults to True.

    Returns:
        HuggingFaceLLM: A configured HuggingFaceLLM instance ready for use with LlamaIndex.

    [all]
    """
    # Log the start of the model and tokenizer loading process
    logging.info("Loading Qwen model and tokenizer...")
    # Load the pre-trained causal language model from Hugging Face
    qwen_model = AutoModelForCausalLM.from_pretrained(
        model_name,
        torch_dtype="auto",
        device_map=device_map,
        trust_remote_code=trust_remote_code,
    )
    # Load the tokenizer corresponding to the pre-trained model
    qwen_tokenizer = AutoTokenizer.from_pretrained(
        model_name,
        trust_remote_code=trust_remote_code,
    )

    # Log the start of building the HuggingFaceLLM wrapper
    logging.info("Building HuggingFaceLLM wrapper...")
    # Construct a HuggingFaceLLM instance with custom generation parameters
    llm = HuggingFaceLLM(
        model_name="Qwen/Qwen2.5-0.5B-Instruct",
        tokenizer_name="Qwen/Qwen2.5-0.5B-Instruct",  # <-- ensure they match
        max_new_tokens=4096,
        context_window=31768,
    )
    # Return the configured HuggingFaceLLM instance
    return llm


# Initialize the language model using the build_qwen_llm function
llm = build_qwen_llm()

# Example of "complete" style usage of the language model
response_text = llm.complete("Tell me about large language models.")
print("LLM Completion:", response_text)

# Example of "chat" style usage of the language model
user_message = ChatMessage(
    role=MessageRole.USER, content="Tell me about large language models."
)
chat_response = llm.chat(messages=[user_message])
print("LLM Chat Response:", chat_response.message.content)

# Load news data from a CSV file
news = pd.read_csv(
    "https://raw.githubusercontent.com/tomasonjo/blog-datasets/main/news_articles.csv"
)[:]

# Display the first few rows of the news DataFrame
news.head()

# ================= Settings =================
# Limit the number of documents to process (default: 50)
MAX_DOCUMENTS = 500

# Define a maximum token limit per document to avoid model sequence length errors
MAX_TOKENS_PER_DOCUMENT = 16000

# ================= Loading Documents =================

# Define the path to the saved documents pickle file
documents_pickle_path = "/documents/extracted_text_output/documents.pkl"

# Log the start of document loading
print("[INFO] Starting document loading...")

# Load the complete document list from the saved file if present
if os.path.exists(documents_pickle_path):
    with open(documents_pickle_path, "rb") as f:
        # Load the documents from the pickle file
        loaded_documents = pickle.load(f)
        # Check if the loaded data is a list
        if not isinstance(loaded_documents, list):
            raise ValueError("Pickle file does not contain a list of documents")
        # Log the number of loaded documents
        print(
            f"[INFO] Loaded {len(loaded_documents)} documents from {documents_pickle_path}"
        )
else:
    # Log an error if the documents file is not found
    print(f"[ERROR] Documents file not found at {documents_pickle_path}")
    # Initialize an empty list if no documents are loaded
    loaded_documents = []

# Assuming `news` is a pandas DataFrame with 'title' and 'text' columns
if "news" in locals() and not news.empty:
    # Log the start of news document creation
    print("[INFO] Creating news documents...")
    # Use tqdm for progress bar during news document creation
    new_documents = [
        Document(text=f"{row['title']}: {row['text']}")
        for i, row in tqdm(
            news.iterrows(), total=len(news), desc="Creating news documents"
        )
    ]
    # Combine loaded documents with newly created news documents
    documents = loaded_documents + new_documents
else:
    # If no news data is available, use only the loaded documents
    documents = loaded_documents

# Limit the number of documents to process based on the MAX_DOCUMENTS setting
documents = documents[:MAX_DOCUMENTS]

# Split documents exceeding the maximum token limit
split_documents = []
for doc in documents:
    # Check if the document text exceeds the maximum token limit
    if len(doc.text.split()) > MAX_TOKENS_PER_DOCUMENT:
        print(f"[INFO] Splitting a large document exceeding token limit.")
        # Initialize a SentenceSplitter with specified chunk size and overlap
        splitter = SentenceSplitter(
            chunk_size=MAX_TOKENS_PER_DOCUMENT, chunk_overlap=1024
        )
        # Split the document into nodes
        split_nodes = splitter.get_nodes_from_documents([doc])
        # Create new Document objects from the split nodes
        for node in split_nodes:
            split_documents.append(
                Document(text=node.get_content(metadata_mode=MetadataMode.LLM))
            )
    else:
        # If the document does not exceed the limit, add it directly
        split_documents.append(doc)

# Update the documents list with the split documents
documents = split_documents

# ================= Sentence Splitting =================

# Log the start of sentence splitting
print("[INFO] Starting sentence splitting...")

if documents:
    # Initialize SentenceSplitter with desired configuration
    splitter = SentenceSplitter(
        chunk_size=4096,  # Define the desired chunk size
        chunk_overlap=512,  # Define the overlap between chunks
    )

    # Use tqdm for progress bar during node generation
    nodes = list(
        tqdm(
            splitter.get_nodes_from_documents(documents),
            total=len(documents),
            desc="Generating nodes",
        )
    )

    # Output the number of nodes generated
    print(f"[INFO] Generated {len(nodes)} nodes.")
else:
    # Log an error if no documents are available for processing
    print("[ERROR] No documents available to process.")

# ================= Configurations =================


class Config:
    """
    Configuration class to hold various settings for the document processing pipeline.

    This class defines constants for supported file extensions, drive mount points,
    default output directories, and other configuration parameters.

    [all]
    """

    SUPPORTED_EXTENSIONS = [".pdf", ".py", ".txt", ".json", ".docx", ".md"]
    DRIVE_MOUNT_POINT = "/content/drive"
    DEFAULT_DRIVE_FOLDER = os.path.join(DRIVE_MOUNT_POINT, "My Drive")
    DEFAULT_OUTPUT_DIR = "/documents/extracted_text_output"
    DOCUMENTS_PICKLE_FILE = "documents.pkl"
    CONTEXT_RANGE = 512


# ================= Utility Functions =================

'''
only for on colab
def mount_drive(mount_point: str = Config.DRIVE_MOUNT_POINT) -> None:
    """Mounts Google Drive at the specified mount point."""
    print(f"[INFO] Mounting Google Drive at {mount_point}...")
    drive.mount(mount_point, force_remount=False)
'''


def deduplicate_documents(docs: List[Document]) -> List[Document]:
    """
    Deduplicate Document objects based on the content hash of their text,
    while preserving order.

    This function calculates the SHA256 hash of each document's text content and
    keeps only the first occurrence of documents with the same hash, thus removing duplicates.

    Args:
        docs (List[Document]): A list of Document objects to deduplicate.

    Returns:
        List[Document]: A new list of Document objects with duplicates removed,
        preserving the original order.

    [all]
    """
    # Initialize a set to store seen content hashes
    seen_hashes = set()
    # Initialize a list to store unique documents
    unique_docs = []
    # Iterate through each document in the input list
    for doc in docs:
        # Calculate the SHA256 hash of the document's text content
        content_hash = hashlib.sha256(doc.text.encode("utf-8")).hexdigest()
        # Check if the content hash has been seen before
        if content_hash not in seen_hashes:
            # If not seen, add the hash to the set and the document to the unique list
            seen_hashes.add(content_hash)
            unique_docs.append(doc)
    # Return the list of unique documents
    return unique_docs


def load_existing_documents(pickle_path: str) -> Tuple[List[Document], set]:
    """
    Load existing documents from a pickle file and create a set of content hashes
    for processed documents. Returns an empty list and set if loading fails.

    This function attempts to load a list of Document objects from a specified pickle file.
    If successful, it also generates a set of SHA256 content hashes for each document
    to track previously processed documents. If loading fails, it returns an empty list
    and an empty set.

    Args:
        pickle_path (str): The path to the pickle file containing the saved documents.

    Returns:
        Tuple[List[Document], set]: A tuple containing:
            - A list of Document objects loaded from the pickle file.
            - A set of SHA256 content hashes for the loaded documents.
            Returns an empty list and set if loading fails.

    [all]
    """
    try:
        # Attempt to open and load the pickle file
        with open(pickle_path, "rb") as f:
            docs = pickle.load(f)
        # Generate a set of content hashes for the loaded documents
        processed_hashes = {
            hashlib.sha256(doc.text.encode("utf-8")).hexdigest() for doc in docs
        }
        # Return the loaded documents and their content hashes
        return docs, processed_hashes
    except Exception:
        # If any exception occurs during loading, return an empty list and set
        return [], set()


# ================= Document Extraction =================


def extract_text_from_file(file, filename: str) -> str:
    """Extracts text from a single file, handling various formats.

    This function reads a file and extracts its text content, handling different file
    formats such as .txt, .py, .json, .pdf, and .docx. It uses appropriate libraries
    to parse each file type and returns the extracted text as a string.

    Args:
        file (file-like object): A file-like object representing the file to be read.
        filename (str): The name of the file, used to determine the file type.

    Returns:
        str: The extracted text content from the file. Returns an empty string if
        extraction fails or if the file type is not supported.

    [all]
    """
    # Initialize an empty string to store the extracted content
    content = ""
    # Convert the filename to lowercase for case-insensitive comparison
    lower_name = filename.lower()
    try:
        # Check if the file is a text-based file
        if lower_name.endswith((".txt", ".py", ".json")):
            # Read the file content and decode it as UTF-8, replacing errors
            content = file.read().decode("utf-8", errors="replace")
        # Check if the file is a PDF
        elif lower_name.endswith(".pdf"):
            # Wrap the file content in BytesIO to ensure binary mode
            pdf_binary = BytesIO(file.read())
            # Create a PdfReader object
            reader = PdfReader(pdf_binary)
            # If PDF is encrypted, try to decrypt
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception as e:
                    print(f"[ERROR] Failed to decrypt {filename}: {e}")
                    return content
            # Extract text from each page of the PDF
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content += text
        # Check if the file is a DOCX
        elif lower_name.endswith(".docx"):
            try:
                # Create a Document object from the DOCX file
                doc = docx.Document(file)
                # Extract text from each paragraph
                for para in doc.paragraphs:
                    content += para.text + "\n"
            except zipfile.BadZipFile:
                # Skip invalid .docx files gracefully
                pass
    except Exception as e:
        # Log an error if any exception occurs during extraction
        print(f"[ERROR] Failed to extract from {filename}: {e}")
    # Return the extracted content
    return content


def process_file(
    filepath: str, supported_exts: List[str], processed_hashes: set
) -> List[Document]:
    """
    Process a single file if its content hasn't been processed before.
    Extract its content if supported, and return a Document object in a list (empty if unsupported).

    This function checks if a file's extension is supported, then extracts its content.
    It skips processing if the file's content hash has already been processed.
    If the file is supported and not a duplicate, it creates a Document object and returns it in a list.

    Args:
        filepath (str): The path to the file to be processed.
        supported_exts (List[str]): A list of supported file extensions.
        processed_hashes (set): A set of content hashes of already processed files.

    Returns:
        List[Document]: A list containing a single Document object if the file is
        processed successfully, or an empty list if the file is unsupported or a duplicate.

    [all]
    """
    # Extract the filename from the filepath
    filename = os.path.basename(filepath)
    # Convert the filename to lowercase for case-insensitive comparison
    lower_name = filename.lower()
    # Check if the file extension is supported
    if any(lower_name.endswith(ext) for ext in supported_exts):
        try:
            # Open the file in binary read mode
            with open(filepath, "rb") as f:
                # Extract text content from the file
                content = extract_text_from_file(f, filename)
            # Check if content was extracted
            if content:
                # Generate a content hash for the file and its content
                content_hash = hashlib.sha256(
                    f"{filename}{content}".encode("utf-8")
                ).hexdigest()
                # Skip if content hash already processed
                if content_hash in processed_hashes:
                    return []
                # Add new hash to set to avoid future duplicates in this run
                processed_hashes.add(content_hash)
                # Create a Document with filename as title
                doc_text = f"{filename}: {content}"
                # Return a list containing the new Document object
                return [Document(text=doc_text)]
        except Exception as e:
            # Log an error if any exception occurs during processing
            print(f"[ERROR] Failed to process {filepath}: {e}")
    # Check if the file is a zip file
    elif lower_name.endswith(".zip"):
        # Handle zip files separately
        return extract_documents_from_zip(filepath, supported_exts)
    # Return an empty list if the file is not supported
    return []


def extract_documents_from_zip(
    zip_filepath: str, supported_exts: List[str]
) -> List[Document]:
    """Recursively extracts Document objects from supported files inside a zip file.

    This function opens a zip file and iterates through its contents, extracting text
    from supported file types. It handles nested zip files recursively.

    Args:
        zip_filepath (str): The path to the zip file.
        supported_exts (List[str]): A list of supported file extensions.

    Returns:
        List[Document]: A list of Document objects extracted from the zip file.

    [all]
    """
    # Initialize an empty list to store extracted documents
    documents = []
    try:
        # Open the zip file in read mode
        with zipfile.ZipFile(zip_filepath, "r") as zip_ref:
            # Iterate through each file in the zip archive
            for file_info in zip_ref.infolist():
                # Extract the filename
                filename = file_info.filename
                # Skip directories
                if file_info.is_dir():
                    continue
                # Convert the filename to lowercase for case-insensitive comparison
                lower_name = filename.lower()
                # Check if the file extension is supported
                if any(lower_name.endswith(ext) for ext in supported_exts):
                    try:
                        # Open the file within the zip archive
                        with zip_ref.open(file_info) as file:
                            # Extract text content from the file
                            content = extract_text_from_file(file, filename)
                            # If content was extracted, create a Document object
                            if content:
                                doc_text = f"{filename}: {content}"
                                documents.append(Document(text=doc_text))
                    except Exception as e:
                        # Log an error if any exception occurs during processing
                        print(f"[ERROR] Error processing file {filename} in zip: {e}")
                # Check if the file is a nested zip file
                elif lower_name.endswith(".zip"):
                    # Create a temporary path for the nested zip file
                    nested_zip_path = os.path.join("/tmp", os.path.basename(filename))
                    # Write the nested zip file to the temporary path
                    with open(nested_zip_path, "wb") as nested_file:
                        nested_file.write(zip_ref.read(file_info))
                    # Recursively extract documents from the nested zip file
                    documents.extend(
                        extract_documents_from_zip(nested_zip_path, supported_exts)
                    )
                    # Remove the temporary nested zip file
                    os.remove(nested_zip_path)
    except zipfile.BadZipFile:
        # Log an error if the zip file is invalid
        print(f"[ERROR] Bad zip file: {zip_filepath}")
    # Return the list of extracted documents
    return documents


def extract_documents_from_drive(
    folder_path: str, supported_exts: List[str], processed_hashes: set
) -> List[Document]:
    """
    Recursively processes all files and folders in a Drive directory,
    returning a list of new Document objects that haven't been processed yet.

    This function traverses a directory, processing each file and subdirectory.
    It uses the `process_file` function to extract text from supported files and
    skips files that have already been processed.

    Args:
        folder_path (str): The path to the directory to process.
        supported_exts (List[str]): A list of supported file extensions.
        processed_hashes (set): A set of content hashes of already processed files.

    Returns:
        List[Document]: A list of Document objects extracted from the directory.

    [all]
    """
    # Initialize an empty list to store extracted documents
    documents = []
    # Traverse the directory tree
    for root, _, files in os.walk(folder_path):
        # Iterate through each file in the current directory
        for file in tqdm(files, desc=f"Processing files in {root}"):
            # Construct the full file path
            file_path = os.path.join(root, file)
            # Process the file and extend the documents list
            docs = process_file(file_path, supported_exts, processed_hashes)
            documents.extend(docs)
    # Deduplicate the documents and return the result
    return deduplicate_documents(documents)


# ================= Saving and Updating Results =================


def save_documents(
    documents: List[Document], output_dir: str = Config.DEFAULT_OUTPUT_DIR
) -> None:
    """
    Save the final list of Document objects to a pickle file for later reloading.
    Updates existing pickle by appending new documents.

    This function saves a list of Document objects to a pickle file. It first loads
    any existing documents from the pickle file, merges them with the new documents,
    deduplicates the combined list, and then saves the updated list back to the pickle file.

    Args:
        documents (List[Document]): A list of Document objects to save.
        output_dir (str): The directory where the pickle file will be saved.
            Defaults to `Config.DEFAULT_OUTPUT_DIR`.

    [all]
    """
    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)
    # Construct the full path to the pickle file
    pickle_path = os.path.join(output_dir, Config.DOCUMENTS_PICKLE_FILE)

    # Load existing documents and their content hashes
    existing_docs, processed_hashes = load_existing_documents(pickle_path)

    # Merge new documents with existing ones
    all_documents = existing_docs + documents
    # Deduplicate the merged list of documents
    unique_documents = deduplicate_documents(all_documents)

    # Save updated unique documents list back to pickle
    with open(pickle_path, "wb") as f:
        pickle.dump(unique_documents, f)
    # Log the number of saved documents and the save path
    print(f"[INFO] Saved {len(unique_documents)} unique documents to: {pickle_path}")


# ================= Main Execution =================
def process() -> List[Document]:
    """
    Main function to process documents from a specified directory.

    This function orchestrates the document processing pipeline, including loading
    previously processed documents, extracting new documents, and saving the updated
    list of documents.

    Returns:
        List[Document]: A list of Document objects that were processed.

    [all]
    """
    # Load previously processed documents to avoid reprocessing
    pickle_path = os.path.join("./documents", Config.DOCUMENTS_PICKLE_FILE)
    _, processed_hashes = load_existing_documents(pickle_path)

    # Extract new documents while skipping duplicates based on content hash
    documents = extract_documents_from_drive(
        "./documents", Config.SUPPORTED_EXTENSIONS, processed_hashes
    )

    # Make extracted documents available globally for later access
    global extracted_documents
    extracted_documents = documents

    # Save updated list of documents
    save_documents(documents, "./documents")
    # Log the completion of the processing
    print("[INFO] Processing complete.")
    # Return the list of processed documents
    return documents


def process_document(document):
    """Process a single document and return nodes with error handling."""
    try:
        logging.debug(f"Processing document: {document}")
        nodes = splitter.get_nodes_from_documents(
            [document]
        )  # Process one document at a time
        logging.debug(f"Generated {len(nodes)} nodes for document: {document}")
        return nodes
    except Exception as e:
        logging.error(
            f"Error processing document: {document}\n{traceback.format_exc()}"
        )
        return []  # Return empty list in case of error


# Assuming 'documents' is a list of Document objects
splitter = SentenceSplitter(chunk_size=4096, chunk_overlap=512)
all_nodes = []

# Dynamically determine the number of workers based on available CPU cores
num_workers = (
    os.cpu_count() or multiprocessing.cpu_count()
)  # Fallback to multiprocessing if os.cpu_count() fails

# Parallel execution using ThreadPoolExecutor with dynamic worker count
with ThreadPoolExecutor(max_workers=num_workers) as executor:
    futures = [executor.submit(process_document, doc) for doc in documents]
    for future in tqdm(
        as_completed(futures), total=len(futures), desc="Processing documents"
    ):  # Progress bar
        all_nodes.extend(future.result())

logging.info(f"Total nodes generated: {len(all_nodes)}")

total_nodes = all_nodes + nodes
nodes = total_nodes

class GraphRAGExtractor(TransformComponent):

    """Extract triples from a graph.

    Uses an LLM and a simple prompt + output parsing to extract paths (i.e. triples) and entity, relation descriptions from text.

    Args:
        llm (LLM):
            The language model to use.
        extract_prompt (Union[str, PromptTemplate]):
            The prompt to use for extracting triples.
        parse_fn (callable):
            A function to parse the output of the language model.
        num_workers (int):
            The number of workers to use for parallel processing.
        max_paths_per_chunk (int):
            The maximum number of paths to extract per chunk.
    """

    llm: LLM
    extract_prompt: PromptTemplate
    parse_fn: Callable
    num_workers: int
    max_paths_per_chunk: int

    def __init__(
        self,
        llm: Optional[LLM] = None,
        extract_prompt: Optional[Union[str, PromptTemplate]] = None,
        parse_fn: Callable = default_parse_triplets_fn,
        max_paths_per_chunk: int = 10,
        num_workers: int = 4,
    ) -> None:
        """Init params."""
        from llama_index.core import Settings

        if isinstance(extract_prompt, str):
            extract_prompt = PromptTemplate(extract_prompt)

        super().__init__(
            llm=llm or Settings.llm,
            extract_prompt=extract_prompt or DEFAULT_KG_TRIPLET_EXTRACT_PROMPT,
            parse_fn=parse_fn,
            num_workers=num_workers,
            max_paths_per_chunk=max_paths_per_chunk,
        )

    @classmethod
    def class_name(cls) -> str:
        return "GraphExtractor"

    def __call__(
        self, nodes: List[BaseNode], show_progress: bool = False, **kwargs: Any
    ) -> List[BaseNode]:
        """Extract triples from nodes."""
        return asyncio.run(
            self.acall(nodes, show_progress=show_progress, **kwargs)
        )

    async def _aextract(self, node: BaseNode) -> BaseNode:
        """Extract triples from a node."""
        assert hasattr(node, "text")

        text = node.get_content(metadata_mode="llm")
        try:
            llm_response = await self.llm.apredict(
                self.extract_prompt,
                text=text,
                max_knowledge_triplets=self.max_paths_per_chunk,
            )
            entities, entities_relationship = self.parse_fn(llm_response)
        except ValueError:
            entities = []
            entities_relationship = []

        existing_nodes = node.metadata.pop(KG_NODES_KEY, [])
        existing_relations = node.metadata.pop(KG_RELATIONS_KEY, [])
        metadata = node.metadata.copy()
        for entity, entity_type, description in entities:
            metadata[
                "entity_description"
            ] = description  # Not used in the current implementation. But will be useful in future work.
            entity_node = EntityNode(
                name=entity, label=entity_type, properties=metadata
            )
            existing_nodes.append(entity_node)

        metadata = node.metadata.copy()
        for triple in entities_relationship:
            subj, rel, obj, description = triple
            subj_node = EntityNode(name=subj, properties=metadata)
            obj_node = EntityNode(name=obj, properties=metadata)
            metadata["relationship_description"] = description
            rel_node = Relation(
                label=rel,
                source_id=subj_node.id,
                target_id=obj_node.id,
                properties=metadata,
            )

            existing_nodes.extend([subj_node, obj_node])
            existing_relations.append(rel_node)

        node.metadata[KG_NODES_KEY] = existing_nodes
        node.metadata[KG_RELATIONS_KEY] = existing_relations
        return node

    async def acall(
        self, nodes: List[BaseNode], show_progress: bool = False, **kwargs: Any
    ) -> List[BaseNode]:
        """Extract triples from nodes async."""
        jobs = []
        for node in nodes:
            jobs.append(self._aextract(node))

        return await run_jobs(
            jobs,
            workers=self.num_workers,
            show_progress=show_progress,
            desc="Extracting paths from text",
        )

class GraphRAGStore(SimplePropertyGraphStore):
    community_summary = {}
    max_cluster_size = 5

    def generate_community_summary(self, text):
        """Generate summary for a given text using an LLM."""
        messages = [
            ChatMessage(
                role="system",
                content=(
                    "You are provided with a set of relationships from a knowledge graph, each represented as "
                    "entity1->entity2->relation->relationship_description. Your task is to create a summary of these "
                    "relationships. The summary should include the names of the entities involved and a concise synthesis "
                    "of the relationship descriptions. The goal is to capture the most critical and relevant details that "
                    "highlight the nature and significance of each relationship. Ensure that the summary is coherent and "
                    "integrates the information in a way that emphasizes the key aspects of the relationships."
                ),
            ),
            ChatMessage(role="user", content=text),
        ]
        response = llm.chat(messages)
        clean_response = re.sub(r"^assistant:\s*", "", str(response)).strip()
        return clean_response

    def build_communities(self):
        """Builds communities from the graph and summarizes them."""
        nx_graph = self._create_nx_graph()
        community_hierarchical_clusters = hierarchical_leiden(
            nx_graph, max_cluster_size=self.max_cluster_size
        )
        community_info = self._collect_community_info(
            nx_graph, community_hierarchical_clusters
        )
        self._summarize_communities(community_info)

    def _create_nx_graph(self):
        """Converts internal graph representation to NetworkX graph."""
        nx_graph = nx.Graph()
        for node in self.graph.nodes.values():
            nx_graph.add_node(str(node))
        for relation in self.graph.relations.values():
            nx_graph.add_edge(
                relation.source_id,
                relation.target_id,
                relationship=relation.label,
                description=relation.properties["relationship_description"],
            )
        return nx_graph

    def _collect_community_info(self, nx_graph, clusters):
        """Collect detailed information for each node based on their community."""
        community_mapping = {item.node: item.cluster for item in clusters}
        community_info = {}
        for item in clusters:
            cluster_id = item.cluster
            node = item.node
            if cluster_id not in community_info:
                community_info[cluster_id] = []

            for neighbor in nx_graph.neighbors(node):
                if community_mapping[neighbor] == cluster_id:
                    edge_data = nx_graph.get_edge_data(node, neighbor)
                    if edge_data:
                        detail = f"{node} -> {neighbor} -> {edge_data['relationship']} -> {edge_data['description']}"
                        community_info[cluster_id].append(detail)
        return community_info

    def _summarize_communities(self, community_info):
        """Generate and store summaries for each community."""
        for community_id, details in community_info.items():
            details_text = (
                "\n".join(details) + "."
            )  # Ensure it ends with a period
            self.community_summary[
                community_id
            ] = self.generate_community_summary(details_text)

    def get_community_summaries(self):
        """Returns the community summaries, building them if not already done."""
        if not self.community_summary:
            self.build_communities()
        return self.community_summary
    
class GraphRAGQueryEngine(CustomQueryEngine):
    graph_store: GraphRAGStore
    llm: LLM

    def custom_query(self, query_str: str) -> str:
        """Process all community summaries to generate answers to a specific query."""
        community_summaries = self.graph_store.get_community_summaries()
        community_answers = [
            self.generate_answer_from_summary(community_summary, query_str)
            for _, community_summary in community_summaries.items()
        ]

        final_answer = self.aggregate_answers(community_answers)
        return final_answer

    def generate_answer_from_summary(self, community_summary, query):
        """Generate an answer from a community summary based on a given query using LLM."""
        prompt = (
            f"Given the community summary: {community_summary}, "
            f"how would you answer the following query? Query: {query}"
        )
        messages = [
            ChatMessage(role="system", content=prompt),
            ChatMessage(
                role="user",
                content="I need an answer based on the above information.",
            ),
        ]
        response = self.llm.chat(messages)
        cleaned_response = re.sub(r"^assistant:\s*", "", str(response)).strip()
        return cleaned_response

    def aggregate_answers(self, community_answers):
        """Aggregate individual community answers into a final, coherent response."""
        # intermediate_text = " ".join(community_answers)
        prompt = "Combine the following intermediate answers into a final, concise response."
        messages = [
            ChatMessage(role="system", content=prompt),
            ChatMessage(
                role="user",
                content=f"Intermediate answers: {community_answers}",
            ),
        ]
        final_response = self.llm.chat(messages)
        cleaned_final_response = re.sub(
            r"^assistant:\s*", "", str(final_response)
        ).strip()
        return cleaned_final_response
    
KG_TRIPLET_EXTRACT_TMPL = """
-Goal-
Given a text document, identify all entities and their entity types from the text and all relationships among the identified entities.
Given the text, extract up to {max_knowledge_triplets} entity-relation triplets.

-Steps-
1. Identify all entities. For each identified entity, extract the following information:
- entity_name: Name of the entity, capitalized
- entity_type: Type of the entity
- entity_description: Comprehensive description of the entity's attributes and activities
Format each entity as ("entity"$$$$"<entity_name>"$$$$"<entity_type>"$$$$"<entity_description>")

2. From the entities identified in step 1, identify all pairs of (source_entity, target_entity) that are *clearly related* to each other.
For each pair of related entities, extract the following information:
- source_entity: name of the source entity, as identified in step 1
- target_entity: name of the target entity, as identified in step 1
- relation: relationship between source_entity and target_entity
- relationship_description: explanation as to why you think the source entity and the target entity are related to each other

Format each relationship as ("relationship"$$$$"<source_entity>"$$$$"<target_entity>"$$$$"<relation>"$$$$"<relationship_description>")

3. When finished, output.

-Real Data-
######################
text: {text}
######################
output:"""
entity_pattern = r'\("entity"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\)'
relationship_pattern = (
    r'\("relationship"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\)'
)

# Define default patterns for entity and relationship extraction, making them configurable
DEFAULT_ENTITY_PATTERN = r'\("entity"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\)'
DEFAULT_RELATIONSHIP_PATTERN = r'\("relationship"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\)'

# loads the base QWEN model for embeddings as well, to keep the system unified
embed_model = HuggingFaceEmbedding(model_name="Qwen/Qwen2.5-0.5B-Instruct")
test_embeds = embed_model.get_text_embedding("Hello World!")
print(test_embeds)
Settings.embed_model = embed_model

from llama_index.core import PropertyGraphIndex

entity_pattern = r'\("entity"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\)'
relationship_pattern = r'\("relationship"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\$\$\$\$"(.+?)"\)'


def parse_fn(response_str: str) -> Any:
    entities = re.findall(entity_pattern, response_str)
    relationships = re.findall(relationship_pattern, response_str)
    return entities, relationships


kg_extractor = GraphRAGExtractor(
    llm=llm,
    extract_prompt=KG_TRIPLET_EXTRACT_TMPL,
    max_paths_per_chunk=2,
    parse_fn=parse_fn,
)
index = PropertyGraphIndex(
    nodes=nodes,
    property_graph_store=GraphRAGStore(),
    kg_extractors=[kg_extractor],
    show_progress=True,
)

list(index.property_graph_store.graph.nodes.values())[-1]
list(index.property_graph_store.graph.relations.values())[0]
list(index.property_graph_store.graph.relations.values())[0].properties[
    "relationship_description"
]

index.property_graph_store.build_communities()
query_engine = GraphRAGQueryEngine(graph_store=index.property_graph_store, llm=llm)
response = query_engine.query("What are the main news discussed in the document?")
display(Markdown(f"{response.response}"))
response = query_engine.query("What are news related to financial sector?")
display(Markdown(f"{response.response}"))
